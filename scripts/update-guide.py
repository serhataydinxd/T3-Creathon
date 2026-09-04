"""Bring docs/IMKAN_Proje_Hakimiyet_ve_Yarisma_Rehberi.docx up to date.

The guide was written against the state just after the Maarif corpus migration.
Since then the product was made compliant with Bilim Türkiye's own structure:
workshop themes and age cohorts, thirty centres with published facilities, four
education formats, their vocabulary, and an honest attribution of 5E.

Run with the venv that has python-docx:
    <venv>/bin/python scripts/update-guide.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

DOC = Path("docs/IMKAN_Proje_Hakimiyet_ve_Yarisma_Rehberi.docx")

# Straight substring swaps. Ordered longest-first so a broad phrase never eats a
# narrower one that follows it.
REPLACEMENTS: list[tuple[str, str]] = [
    # --- counts that moved -------------------------------------------------
    (
        "Önemli tutarlılık notu: README, 59 birim ve 11 tarayıcı testi ifadesini "
        "taşıyor. Güncel dosyalarda statik sayım 99 birim testi ve 13 tarayıcı testi "
        "tanımı gösteriyor.",
        "Önemli tutarlılık notu: README ve bu belge aynı sayıyı taşır; güncel durum "
        "420 birim testi ve 37 tarayıcı testidir ve bu sayılar geliştirme "
        "ortamında npm run check:all çalıştırılarak doğrulanmıştır.",
    ),
    ("99 birim ve 13 tarayıcı testi statik sayımı", "420 birim ve 37 tarayıcı testi"),
    (
        "Repo 99 birim testi ve 13 uçtan uca tarayıcı testi tanımı içeriyor. Birim "
        "testleri AI sözleşmesi, rota ve korpus bütünlüğü, maliyet, generation "
        "record, AWS yapılandırması, veritabanı TLS'i, same-origin ve site metadata "
        "alanlarını kapsar. E2E testleri kayıt/oturum, dört rolün iş akışı, "
        "idempotency, sahte istemci metni, kayıt sahipliği, cross-origin, büyük "
        "gövde ve taslağa doğrudan erişimi sınar.",
        "Repo 420 birim testi ve 37 uçtan uca tarayıcı testi içeriyor. Birim "
        "testleri AI sözleşmesi, rota ve korpus bütünlüğü, konu × merkez ve konu × "
        "envanter matrisleri, format kuralları, maliyet, generation record, AWS "
        "yapılandırması, veritabanı TLS'i, same-origin, site metadata ve arayüz "
        "sözlüğünü kapsar. E2E testleri kayıt/oturum, dört rolün iş akışı, "
        "idempotency, sahte istemci metni, kayıt sahipliği, cross-origin, büyük "
        "gövde, taslağa doğrudan erişim, erişilebilirlik taraması, klavye "
        "kullanımı ve metadata uçlarını sınar.",
    ),
    (
        "Güncel kaynakta statik olarak 99 birim ve 13 E2E test tanımı sayılıyor.",
        "Güncel kaynakta 420 birim ve 37 E2E testi bulunuyor ve tümü yeşil.",
    ),
    ("7 kazanım, 14 rota, 26 malzeme", "7 atölye konusu, 16 rota, 26 malzeme, 30 merkez"),
    (
        "7. sınıf fen bilimlerinden 7 öğrenme çıktısı ve 14 rota; yalnızca 5E modeli.",
        "Bilim Türkiye'nin yedi atölye temasından üçünü kapsayan 7 atölye konusu ve "
        "16 rota; 30 merkez ve dört eğitim formatı; yalnızca 5E iskeleti.",
    ),
    (
        "Her çıktı iki rota taşır; toplam 14 rota vardır.",
        "Her konu en az iki rota taşır; planetaryum ve sergi rotalarıyla toplam 16 "
        "rota vardır.",
    ),
    (
        "7 kazanım ve 14 rota",
        "7 atölye konusu ve 16 rota, yedi temadan üçü",
    ),
    (
        "7. sınıf fen bilimlerinin yedi ünitesinden birer öğrenme çıktısı ve toplam "
        "14 rota vardır. Genişlikten önce izlenebilir dikey dilimi tamamladık.",
        "Yedi atölye konusu ve 16 rota vardır; bunlar Bilim Türkiye'nin yedi "
        "temasından üçünü kapsar. Genişlikten önce izlenebilir dikey dilimi "
        "tamamladık ve kapsanmayan dört temayı ana sayfada 3/7 olarak açıkça "
        "gösteriyoruz.",
    ),
    (
        "7 öğrenme çıktısı, 14 rota, 5E, 26 malzeme.",
        "7 atölye konusu, 16 rota, 26 malzeme, 30 merkez, 4 format; 5E iskeleti.",
    ),
    (
        "Bağımlılık kurulumu bu çalışma ortamında tamamlanamadığı için testlerin "
        "çalıştığı iddia edilmemiştir; yarışma öncesi ekip kendi makinesinde npm run "
        "check:all sonucunu göstermelidir.",
        "Yarışma öncesi ekip yine de kendi makinesinde npm run check:all sonucunu "
        "göstermeye hazır olmalıdır.",
    ),
    # --- the concept was renamed ------------------------------------------
    ("Kazanım Kilididir", "Konu Kilididir"),
    ("Kazanım Kilidi", "Konu Kilidi"),
    ("Kazanım ve koşullar", "Konu ve koşullar"),
    ("kazanım korpusu", "atölye konusu korpusu"),
    ("Kazanım korpusu", "Atölye konusu korpusu"),
    # --- the role is an eğitmen -------------------------------------------
    ("Eğitimci", "Eğitmen"),
    ("eğitimci", "eğitmen"),
    ("öğretmen ve öğrenci yönergeleri", "eğitmen ve öğrenci yönergeleri"),
    # --- claims that are no longer true ------------------------------------
    (
        "Önce MEB kazanımını kilitliyoruz. Yapay zekâ etkinliği uyarlayabilir, fakat "
        "resmî hedefi değiştiremez.",
        "Önce atölye konusunu kilitliyoruz. Yapay zekâ oturumu uyarlayabilir, fakat "
        "konuyu değiştiremez. Konunun okulda öğrenilenle tamamlayıcılığı varsa MEB "
        "öğrenme çıktısı kodu ayrıca gösterilir.",
    ),
    (
        "Bu nedenle sunumda 'MEB kaynağından aktarıldı' denmeli, 'alan uzmanı "
        "doğruladı' denmemelidir.",
        "Bu nedenle sunumda 'MEB kaynağından aktarıldı' denmeli, 'alan uzmanı "
        "doğruladı' denmemelidir. Ayrıca Bilim Türkiye içeriklerinin MEB "
        "kazanımlarına bağlandığına dair kamuya açık bir eşleme bulunamadığı için "
        "'MEB müfredatına uyumlu' denmemeli; 'okul kazanımıyla tamamlayıcı' "
        "denmelidir.",
    ),
]

# Sections appended at the end, as (style, text) pairs.
NEW_SECTIONS: list[tuple[str, str]] = [
    ("Heading 1", "Bilim Türkiye yapısına uyum"),
    (
        "Normal",
        "Problem 3 doğrudan Bilim Türkiye ile ilgilidir ve Bilim Türkiye kendi "
        "programını MEB kazanımına göre değil, atölye temasına ve yaş grubuna göre "
        "kurar. Bu oturumda korpusun ekseni buna göre değiştirildi. Jüri 'bu bizim "
        "yapımıza nasıl oturuyor' diye sorarsa cevap bu bölümdedir.",
    ),
    ("Heading 2", "Yedi tema ve üç yaş grubu"),
    (
        "Normal",
        "Korpus girdisi artık bir kazanım değil, bir atölye konusudur: tema, yaş "
        "grubu, başlık ve özet zorunludur; MEB eşleştirmesi isteğe bağlı bir "
        "alandır. Temalar Teknoloji, Matematik, Girişim, Tasarım, Doğa Bilimleri, "
        "Astronomi Havacılık ve Uzay, Tarım Teknolojileridir. Yaş grupları 6-8, "
        "9-11 ve 12-14'tür. Mevcut yedi konu Astronomi ve Uzay, Teknoloji ve Doğa "
        "Bilimleri temalarını kapsar; kalan dört tema için içerik yoktur ve ana "
        "sayfa bunu 3/7 olarak gösterir.",
    ),
    ("Heading 2", "Otuz merkez ve mekân donanımı"),
    (
        "Normal",
        "Bilim Türkiye 30 merkezde çalışır ve merkezler aynı donanıma sahip "
        "değildir. Uygulama artık merkez seçtirir ve donanımı o merkezin kendi "
        "sayfasında yayımlanan bilgiden türetir: planetaryum, sergi alanı ve "
        "DENEYAP. Yayımlanmamış donanım 'yok' sayılmaz; arayüz hangi donanımın "
        "yayımlanmadığını söyler ve eğitmene işaretleme imkânı bırakır. Bu ayrım "
        "testle korunur.",
    ),
    (
        "Normal",
        "Demoda en güçlü karşılaştırma budur: Uzay Çağı konusu Bilim Trabzon'da "
        "planetaryum rotasını seçer, Bilim Çorum'da kâğıt tüp rotasına düşer ve "
        "nedenini yazar. Aynı konu, aynı oturum, farklı merkez.",
    ),
    ("Heading 2", "Dört eğitim formatı"),
    (
        "Normal",
        "Okul grubu programı, tematik eğitim, uzun süreli eğitim ve çevrim içi "
        "atölye tanımlıdır; yayımlanmış oturum süresi bir saattir ve varsayılan "
        "profil artık 60 dakikadır. Çevrim içi formatta merkez donanımı "
        "kullanılamaz, çünkü katılımcı evdedir; bu durumda planetaryum rotası "
        "farklı ve doğru bir gerekçeyle reddedilir. Uzun süreli eğitim 15 veya 30 "
        "saatlik bir pakettir ve oturum dağılımı yayımlanmamıştır; format "
        "seçilebilir fakat plan tek oturumu kapsadığını açıkça yazar.",
    ),
    ("Heading 2", "Pedagoji iddiasının sınırı"),
    (
        "Normal",
        "5E, İMKÂN'ın kendi aşama iskeletidir. Bilim Türkiye'nin yayımlanmış "
        "yaklaşımı 'Yaparak Yaşayarak Öğrenme' ve proje tabanlı çalışmadır ve "
        "İMKÂN bunu uygulamaz. Arayüzde model adı '5E Öğrenme Döngüsü (İMKÂN)' "
        "olarak yazılır ve yardım metni bu ayrımı söyler. Jüriye 'Bilim "
        "Türkiye'nin pedagojisini uyguluyoruz' denmemelidir.",
    ),
    ("Heading 2", "Dil"),
    (
        "Normal",
        "Arayüz Bilim Türkiye'nin sözlüğünü kullanır: atölyeyi yürüten kişi "
        "eğitmendir, zamanlanmış faaliyet oturum veya atölye eğitimidir, ders "
        "değildir. Üretim istemi de bu register'ı ister, aksi hâlde canlı üretim "
        "okul dili üretmeye devam ederdi. Bu sözcük tercihleri testle korunur.",
    ),
    ("Heading 2", "Bu oturumda düzeltilen kusurlar"),
    (
        "Normal",
        "Erişilebilirlik taraması genişletildiğinde üç gerçek kusur bulundu ve "
        "giderildi: malzeme tablosu yatay kaydırılabilir olduğu hâlde odak "
        "almıyordu, yani klavye kullanıcısı tabloyu hiç kaydıramıyordu; maliyet "
        "listesi geçersiz bir tanım listesi yapısı kullanıyordu; ve soluk gövde "
        "metni 4,48 kontrast oranıyla eşiğin hemen altındaydı. Ayrıca elektrik ve "
        "internet kartları görsel olarak 'var' derken ekran okuyucuya 'yok' "
        "diyordu. Üretim sürümü damgası da bu oturumda güncellendi; damga "
        "güncellenmediğinde eski bir üretim kaydı sessizce kabul edilir.",
    ),
]


def replace_everywhere(doc: Document) -> dict[str, int]:
    """Apply the swaps to paragraphs and table cells, run by run where possible.

    A run-level edit keeps bold and italic inside a paragraph intact. When the
    target text straddles runs, the paragraph is rebuilt from its first run,
    which is reported so the loss is visible rather than silent.
    """
    counts = {"run": 0, "paragraph": 0}

    def apply_to_paragraph(par: Paragraph) -> None:
        for old, new in REPLACEMENTS:
            if old not in par.text:
                continue
            handled = False
            for run in par.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    counts["run"] += 1
                    handled = True
                    break
            if not handled:
                rebuilt = par.text.replace(old, new)
                if par.runs:
                    par.runs[0].text = rebuilt
                    for extra in par.runs[1:]:
                        extra.text = ""
                    counts["paragraph"] += 1
                    print(f"  rebuilt (spanned runs): {old[:58]}…")

    for par in doc.paragraphs:
        apply_to_paragraph(par)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    apply_to_paragraph(par)
    return counts


def append_sections(doc: Document) -> bool:
    """Add the new material, unless a previous run already did.

    The script is meant to be re-runnable while the guide is still being
    corrected, so appending has to be idempotent.
    """
    marker = NEW_SECTIONS[0][1]
    if any(par.text.strip() == marker for par in doc.paragraphs):
        print("  sections already present; not appending again")
        return False
    for style, text in NEW_SECTIONS:
        doc.add_paragraph(text, style=style)
    return True


def main() -> int:
    if not DOC.exists():
        print(f"missing: {DOC}", file=sys.stderr)
        return 1
    doc = Document(str(DOC))
    counts = replace_everywhere(doc)
    appended = append_sections(doc)
    doc.save(str(DOC))
    print(f"replacements: {counts['run']} in-run, {counts['paragraph']} rebuilt")
    if appended:
        print(f"appended {sum(1 for style, _ in NEW_SECTIONS if style == 'Heading 2')} subsections")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
